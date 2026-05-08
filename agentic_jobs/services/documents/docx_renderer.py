from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Literal

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from markdown_it import MarkdownIt

from agentic_jobs.services.documents.style import get_document_style


_MD = MarkdownIt("commonmark").enable("strikethrough")


@dataclass(slots=True)
class DocBlock:
    kind: Literal["paragraph", "heading", "bullet", "ordered"]
    text: str


def render_cover_letter_docx(text: str, output_path: Path) -> Path:
    style = get_document_style()
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(style.margin_top_in)
    section.bottom_margin = Inches(style.margin_bottom_in)
    section.left_margin = Inches(style.margin_left_in)
    section.right_margin = Inches(style.margin_right_in)

    normal = document.styles["Normal"]
    normal.font.name = style.docx_font_name
    normal.font.size = Pt(style.font_size)
    _set_style_font(normal, style.docx_font_name)

    blocks = _extract_blocks(text)
    for block in blocks:
        if not block.text.strip():
            continue
        if block.kind == "bullet":
            para = document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "ordered":
            para = document.add_paragraph(block.text, style="List Number")
        else:
            para = document.add_paragraph(block.text)
        para_format = para.paragraph_format
        para_format.line_spacing = style.line_spacing
        para_format.space_after = Pt(0)
        para_format.space_before = Pt(0)
        if block.kind == "heading":
            for run in para.runs:
                run.bold = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _set_style_font(style, font_name: str) -> None:
    try:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _extract_blocks(text: str) -> list[DocBlock]:
    tokens = _MD.parse(text or "")
    blocks: list[DocBlock] = []
    idx = 0
    while idx < len(tokens):
        token = tokens[idx]
        if token.type == "heading_open":
            inline = tokens[idx + 1]
            blocks.append(DocBlock("heading", _collect_inline_text(inline.children or []).strip()))
            idx += 3
            continue
        if token.type == "paragraph_open":
            inline = tokens[idx + 1]
            blocks.append(DocBlock("paragraph", _collect_inline_text(inline.children or []).strip()))
            idx += 3
            continue
        if token.type in {"bullet_list_open", "ordered_list_open"}:
            ordered = token.type == "ordered_list_open"
            idx = _consume_list_blocks(tokens, idx + 1, blocks, ordered)
            continue
        idx += 1
    return blocks


def _consume_list_blocks(tokens: list, idx: int, blocks: list[DocBlock], ordered: bool) -> int:
    closing_type = "ordered_list_close" if ordered else "bullet_list_close"
    while idx < len(tokens):
        token = tokens[idx]
        if token.type == "list_item_open":
            idx += 1
            parts: list[str] = []
            while idx < len(tokens) and tokens[idx].type != "list_item_close":
                if tokens[idx].type == "paragraph_open":
                    inline = tokens[idx + 1]
                    parts.append(_collect_inline_text(inline.children or []))
                    idx += 3
                else:
                    idx += 1
            idx += 1  # consume list_item_close
            text = " ".join(part.strip() for part in parts if part.strip())
            blocks.append(DocBlock("ordered" if ordered else "bullet", text))
            continue
        if token.type == closing_type:
            return idx + 1
        idx += 1
    return idx


def _collect_inline_text(tokens: Iterable) -> str:
    pieces: list[str] = []
    for token in tokens:
        ttype = token.type
        if ttype == "text":
            pieces.append(token.content)
        elif ttype == "code_inline":
            pieces.append(token.content)
        elif ttype in ("softbreak", "hardbreak"):
            pieces.append(" ")
        elif getattr(token, "children", None):
            pieces.append(_collect_inline_text(token.children))
    return "".join(pieces).strip()
